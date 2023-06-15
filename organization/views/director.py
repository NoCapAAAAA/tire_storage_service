from core import models as m
from django.contrib import messages
from django.shortcuts import render
from docx.shared import Cm
from service.filters import UsersFilterDirector
from service.decorators import group_required
from service.charts import months, colorPrimary, get_year_dict
from django.contrib.auth.models import Group
from organization import forms as f
from django.utils.decorators import method_decorator
from django.views.generic import UpdateView, ListView, TemplateView, CreateView, DetailView
from django.urls import reverse_lazy
from django.http import JsonResponse
from django.utils import timezone
from django.db.models import Count, Sum
from django.db.models.functions import ExtractYear, ExtractMonth
from django.contrib.auth import get_user_model
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from datetime import datetime
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
User = get_user_model()


class DirectorHomeView(TemplateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/home_director.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        last_week = timezone.now() - timezone.timedelta(days=7)
        current_datetime = timezone.now()
        twenty_four_hours_ago = current_datetime - timezone.timedelta(hours=24)
        context['calls'] = m.CallApplication.objects.filter(created_at__gte=twenty_four_hours_ago).count()
        context['orders'] = m.OrderStorage.objects.count()
        context['lastweakusers'] = User.objects.filter(date_joined__gte=last_week).count()
        context['lastweakorders'] = m.OrderStorage.objects.filter(created_at__gte=last_week).count()
        context['instorageorders'] = m.OrderStorage.objects.filter(status=m.OrderStatus.STORAGE).count()
        context['finishorders'] = m.OrderStorage.objects.filter(status=m.OrderStatus.FINISH).count()
        last_month = timezone.now() - timezone.timedelta(days=30)
        address_orders = m.OrderStorage.objects.values('address__address').annotate(total_orders=Count('id')).order_by('-total_orders')
        most_orders_address = address_orders.first()
        context['most_orders_address'] = most_orders_address['address__address'] if most_orders_address else None
        context['total_orders'] = most_orders_address['total_orders'] if most_orders_address else 0

        context['totalpriceorderlastmonth'] = m.OrderStorage.objects.filter(
            is_payed=True,
            status=m.OrderStatus.FINISH,
            created_at__gte=last_month
        ).aggregate(Sum('price'))['price__sum']
        context['totalpriceorderyear'] = m.OrderStorage.objects.filter(is_payed=True, status=m.OrderStatus.FINISH).aggregate(Sum('price'))['price__sum']
        last_year = timezone.now() - timezone.timedelta(days=365)
        context['totalpriceorderlastyear'] = m.OrderStorage.objects.filter(
            is_payed=True,
            status=m.OrderStatus.FINISH,
            created_at__year=last_year.year - 1
        ).aggregate(Sum('price'))['price__sum']
        context['callaplication'] = m.CallApplication.objects.order_by('-pk')[:5]
        return context


class DirectorUsersListView(ListView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/users_list_director.html'
    model = User
    paginate_by = 7

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['users'] = User.objects.all()
        context['filter'] = UsersFilterDirector(self.request.GET, queryset=self.get_queryset())
        return context

    def get_queryset(self, **kwargs):
        search_results = UsersFilterDirector(self.request.GET, self.queryset)
        self.no_search_result = True if not search_results.qs else False
        return search_results.qs.distinct()


class DirectorUpdateUserView(UpdateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/user_get_permissions.html'
    model = User
    form_class = f.SettingsProfile

    def get_queryset(self):
        return User.objects.filter(pk=self.kwargs['pk'])

    def form_valid(self, form):
        messages.success(self.request, "The user was updated successfully.")
        return super(DirectorUpdateUserView, self).form_valid(form)

    def get_success_url(self):
        return reverse_lazy('director_get_user_permissions', kwargs={'pk': self.kwargs['pk']})


class DirectorEmployeeDetailView(UpdateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/detail_employee_director.html'
    model = User
    form_class = f.SettingsProfile

    def form_valid(self, form):
        user = form.save(commit=False)
        user.groups.clear()
        groups = self.request.POST.getlist('groups')
        for group_id in groups:
            group = Group.objects.get(id=group_id)
            user.groups.add(group)
        return super().form_valid(form)

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        group_id = request.POST.get('group_id')
        if group_id:
            group = Group.objects.get(id=group_id)
            self.object.groups.remove(group)
        return super().post(request, *args, **kwargs)

    def get_queryset(self):
        return User.objects.filter(pk=self.kwargs['pk'])

    def get_success_url(self):
        return reverse_lazy('director_detail_employee_view', kwargs={'pk': self.kwargs['pk']})


class DirectorListEmployeeView(ListView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/list_employee_director.html'
    model = User
    paginate_by = 7
    context_object_name = 'employees'
    queryset = User.objects.filter(groups__name='Менеджер')


class DirectorListCallApplicationView(ListView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/list_call_app_director.html'
    model = m.CallApplication
    paginate_by = 7
    context_object_name = 'call_app'


class DirectorDetailCallApplicationView(DetailView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/detail_call_director.html'
    model = m.CallApplication
    context_object_name = 'call'


class DirectorListAddressView(ListView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/address_list.html'
    model = m.AddressService
    paginate_by = 7
    context_object_name = 'address'


class DirectorCreateAddressView(CreateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/address_create.html'
    model = m.AddressService
    fields = ('address',)
    def get_success_url(self):
        return reverse_lazy('director_list_address')


class DirectorEditAddressView(UpdateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/address_edit.html'
    model = m.AddressService
    fields = ('address',)
    def get_success_url(self):
        return reverse_lazy('director_list_address')


class DirectorListPeriodView(ListView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/period_list.html'
    model = m.PeriodOfStorage
    paginate_by = 7
    context_object_name = 'period'


class DirectorCreatePeriodView(CreateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/period_create.html'
    model = m.PeriodOfStorage
    fields = ('period',)
    def get_success_url(self):
        return reverse_lazy('director_list_period')


class DirectorEditPeriodView(UpdateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/period_edit.html'
    model = m.PeriodOfStorage
    fields = ('period',)
    def get_success_url(self):
        return reverse_lazy('director_list_period')


class DirectorListQuantityView(ListView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/quantity_list.html'
    model = m.QuantityOfTires
    paginate_by = 7
    context_object_name = 'quantity'


class DirectorCreateQuantityView(CreateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/quantity_create.html'
    model = m.QuantityOfTires
    fields = ('quantity',)
    def get_success_url(self):
        return reverse_lazy('director_list_quantity')


class DirectorEditQuantityView(UpdateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/quantity_edit.html'
    model = m.QuantityOfTires
    fields = ('quantity',)
    def get_success_url(self):
        return reverse_lazy('director_list_quantity')


class DirectorListSizeView(ListView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/tire_size_list.html'
    model = m.TireSize
    paginate_by = 7
    context_object_name = 'size'


class DirectorCreateSizeView(CreateView):
    #
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/tire_size_create.html'
    model = m.TireSize
    fields = ('size',)
    def get_success_url(self):
        return reverse_lazy('director_list_size')




class DirectorEditSizeView(UpdateView):
    @method_decorator(group_required('Директор'))
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    template_name = 'director/references/tire_size_edit.html'
    model = m.TireSize
    fields = ('size',)
    def get_success_url(self):
        return reverse_lazy('director_list_size')


#################
# ВСЁ ОСТАЛЬНОЕ
#################

def get_filter_options(request):
    grouped_purchases = m.OrderStorage.objects.annotate(year=ExtractYear("created_at")).values("year").order_by("-year").distinct()
    options = [purchase["year"] for purchase in grouped_purchases]

    return JsonResponse({
        "options": options,
    })


def get_sales_chart(request, year):
    purchases = m.OrderStorage.objects.filter(created_at__year=year)
    grouped_purchases = purchases.annotate(month=ExtractMonth("created_at"))\
        .values("month").annotate(count=Count("id")).values("month", "count").order_by("month")

    sales_dict = get_year_dict()

    for group in grouped_purchases:
        sales_dict[months[group["month"]-1]] = group["count"]

    return JsonResponse({
        "title": f"Sales in {year}",
        "data": {
            "labels": list(sales_dict.keys()),
            "datasets": [{
                "label": "Количество заказов",
                "backgroundColor": colorPrimary,
                "borderColor": colorPrimary,
                "data": list(sales_dict.values()),
            }]
        },
    })


def generate_report(request):
    if request.method == 'POST':
        date_from = request.POST.get('date_from')
        date_until = request.POST.get('date_until')

        # Получение заказов в выбранном диапазоне дат
        orders = m.OrderStorage.objects.filter(created_at__range=[date_from, date_until]).order_by('created_at')

        # Создание нового документа
        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Вставка картинки в ячейку таблицы
        cell = table.cell(0, 0)
        image_path = 'C:\\Users\\user\\Desktop\\каретников_диплом\\scr\\logo.png'
        cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2), height=Inches(1))
        document.add_paragraph(f'                                                                                                                 ООО ХРАНЕНИЕ ШИН')
        document.add_paragraph(f'                                                                                                                 ИНН 928131283112')
        document.add_paragraph(f'                                                                                                                 ОГРНИП 928134567831283112')
        # Добавление заголовка
        date_from_formatted = datetime.strptime(date_from, "%Y-%m-%d").strftime("%d-%m-%Y")
        date_until_formatted = datetime.strptime(date_until, "%Y-%m-%d").strftime("%d-%m-%Y")

        document.add_heading(f'Отчет о заказах с {date_from_formatted} по {date_until_formatted}', level=1)

        # Создание таблицы и добавление заголовков столбцов
        table = document.add_table(rows=1, cols=8)
        table.style = 'Table Grid'

        # Установка размера таблицы
        table.width = Cm(25)  # Установите ширину таблицы в сантиметрах
        table.height = Cm(10)  # Установите высоту таблицы в сантиметрах

        headers = ['Номер заказа', 'Клиент', 'Размер шин', "Период", "Адрес сервиса", 'Статус заказа', 'Создан', 'Стоимость заказа']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        total_cost = 0
        # Заполнение таблицы данными
        for i, order in enumerate(orders):
            row_cells = table.add_row().cells
            row_cells[0].text = str(order.pk)
            row_cells[1].text = str(order.user.first_name)
            row_cells[2].text = str(order.size) + ' R'
            row_cells[3].text = str(order.period) + ' мес.'
            row_cells[4].text = str(order.address)
            row_cells[5].text = str(order.get_status_display())
            row_cells[6].text = str(order.created_at.strftime ("%Y-%m-%d"))
            row_cells[7].text = str(order.price) + ' руб.'
            total_cost += order.price
        document.add_paragraph(f'')
        document.add_paragraph(f'')
        document.add_paragraph(f'')
        document.add_paragraph(f'                                                                                                                 Итого: {total_cost} руб.')


        # Создание HTTP-ответа для скачивания документа
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=report.docx'

        # Сохранение документа в HTTP-ответе
        document.save(response)

        return response

    return render(request, 'director/report.html')


def generate_report_users(request):
    if request.method == 'POST':
        date_from = request.POST.get('date_from')
        date_until = request.POST.get('date_until')

        # Получение заказов в выбранном диапазоне дат
        users = User.objects.filter(date_joined__range=[date_from, date_until]).order_by('date_joined')

        # Создание нового документа
        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Вставка картинки в ячейку таблицы
        cell = table.cell(0, 0)
        image_path = 'C:\\Users\\user\\Desktop\\каретников_диплом\\scr\\logo.png'
        cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2), height=Inches(1))
        document.add_paragraph(f'                                                                                                                 ООО ХРАНЕНИЕ ШИН')
        document.add_paragraph(f'                                                                                                                 ИНН 928131283112')
        document.add_paragraph(f'                                                                                                                 ОГРНИП 928134567831283112')
        # Добавление заголовка


        date_from_formatted = datetime.strptime(date_from, "%Y-%m-%d").strftime("%d-%m-%Y")
        date_until_formatted = datetime.strptime(date_until, "%Y-%m-%d").strftime("%d-%m-%Y")

        document.add_heading(f'Отчет о пользователях с {date_from_formatted} по {date_until_formatted}', level=1)


        # Создание таблицы и добавление заголовков столбцов
        table = document.add_table(rows=1, cols=7)
        table.style = 'Table Grid'

        # Установка размера таблицы
        table.width = Cm(25)  # Установите ширину таблицы в сантиметрах
        table.height = Cm(10)  # Установите высоту таблицы в сантиметрах

        headers = ['ID', 'Фамилия', 'Имя', "Отчество", "Номер телефона", 'Email', 'Дата присоединения']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        total_cost = users.count()
        # Заполнение таблицы данными
        for i, user in enumerate(users):
            row_cells = table.add_row().cells
            row_cells[0].text = str(user.pk)
            row_cells[1].text = str(user.first_name)
            row_cells[2].text = str(user.last_name)
            row_cells[3].text = str(user.middle_name)
            row_cells[4].text = str(user.phone_number)
            row_cells[5].text = str(user.email)
            row_cells[6].text = str(user.date_joined.strftime("%d-%m-%Y"))

        document.add_paragraph(f'')
        document.add_paragraph(f'')
        document.add_paragraph(f'')
        document.add_paragraph(f'                                                                                            Итого: {total_cost} новых пользотвалей')


        # Создание HTTP-ответа для скачивания документа
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=report.docx'

        # Сохранение документа в HTTP-ответе
        document.save(response)

        return response

    return render(request, 'director/report_users.html')